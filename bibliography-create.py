import docx
from docx.shared import RGBColor
import re

RE_LINESTART = re.compile(r'^\d{2} *')
bibl = docx.Document(r'./sample_files/litspisok.docx')

RE_REPLASE = re.compile(r'\[([0-9,\s]+)\]')

origin_list = []
cell_list = []
replace_el = []
# количество абзацев в документе
# print(len(bibl.paragraphs))
# количество таблиц в файле
# print(len(bibl.tables))

# парсим текст параграфов
for par in bibl.paragraphs:
    print('Количество параграфов', par.text)

table = bibl.tables[0].rows
for row in table:
    cell_list = []
    for cell in row.cells:
        cell_list.append(cell.text)

    origin_list.append(cell_list)
# print(origin_list)

#
# for el in origin_list:
#     print(len(el))
#     print([el[0]])
#     print([el[1]])
#     break
# сортировка по второму эдементу списка
# origin_list.sort(key=lambda x: x[0])
# print(origin_list)
# создание сортированого списка.
# for inx, el in enumerate(origin_list, 1):
# el.append(inx)
# print(origin_list)
# создание словаря со значения бывшего и нового порядка
# origins_dict = {y: [x, z] for x, y, z in tuple(origin_list)}

# print(origins_dict)

# текст первого абзаца в документе
# print(oks_doc12.paragraphs[0].text)

# текст второго абзаца в документе
# print(oks_doc12.paragraphs[6].text)
# print(oks_doc12.paragraphs[7].text)
# print(oks_doc12.paragraphs[8].text)
# print(oks_doc12.paragraphs[9].text)
# текст первого Run второго абзаца


new_bibl = docx.Document()
# final_list = [[x,y] for x in origins_dict.values() for y in origins_dict.keys()]
table = new_bibl.add_table(rows=1, cols=3)
table.style = 'Table Grid'
table.rows[0].cells[0].text = "New"
table.rows[0].cells[1].text = 'Old'
table.rows[0].cells[2].text = "Text"
# print (final_list)
# new_bibl.add_paragraph(final_list)
# origin_list.sort()
sorted_bibl_list = sorted(origin_list, key=lambda x: x[1])
# print(sorted_bibl_list)
for i in sorted_bibl_list:
    try:
        if i[1][0] == 'А':
            # print('найден индекс', sorted_bibl_list.index(i))
            rus_firs_list = sorted_bibl_list[sorted_bibl_list.index(i):]
            rus_firs_list.extend(sorted_bibl_list[:sorted_bibl_list.index(i)])
            break
    except IndexError:
        pass

# добавление новоно индекса в список
for inx, el in enumerate(rus_firs_list, 1):
    el.append(inx)

# создание словаря из списка
resul_dict = {old: [new, text] for new, text, old in rus_firs_list}

# print(resul_dict.keys())

# for i, el in rus_firs_list:
#     row_cells = table.add_row().cells
#     row_cells[0].text = i
#     row_cells[1].text = el


for i in resul_dict:
    row_cells = table.add_row().cells
    row_cells[0].text = str(i)  # не хочет работать с цифрой пришлось сделать строку
    row_cells[1].text = str(resul_dict[i][0])
    row_cells[2].text = resul_dict[i][1]

# for i, el1 in (sorted_bibl_list):
#     row_cells = table.add_row().cells
#     row_cells[0].text = str(i)
#     row_cells[1].text = el1
# запись сортированного списка в фаил
# for i, el1 in enumerate(rus_firs_list, 1):
#
#     if el1[0] != '' and el1[1] != '':
#         row_cells = table.add_row().cells
#         row_cells[0].text = str(i)
#         row_cells[1].text = el1[0]
#         row_cells[2].text = el1[1]
#         row_cells[2].text = el2


new_bibl.save(r'./sample_files/new_bibl.docx')

# открываем новый документ
new_dis_doc = docx.Document(r'./sample_files/dis-try.docx')
# назначени цвета тексту без docx.shared не работает
# runn = new_dis_doc.add_paragraph().add_run('fdffdsfsdfsdf')
# runn.font.color.rgb = RGBColor(255, 0, 0)


for el in new_dis_doc.paragraphs:
    it = re.finditer(RE_REPLASE, el.text)
    for m in it:
        # print(el.text)
        # print(m)
        # print(m.span())
        # print(m.group())

        # создание списка найденых номеров в строке
        find_el = m.group()[1:-1].replace(' ', '').split(',')
        # перебираем и заменяем найденые номера

        replace_el = []
        for old in find_el:
            # print('старый', old, 'новый', resul_dict[int(old)][0])
            replace_el.append(resul_dict[int(old)][0])

        print('старый текст', el.text[m.span()[0] + 1:m.span()[1] - 1])

        # формируем строку для вставки
        input_str = '[{}]'.format(', '.join(replace_el))

        el.text = re.sub(RE_REPLASE, input_str, el.text)
        print('новый текст', el.text[m.span()[0] + 1:m.span()[1] - 1])

new_dis_doc.save(r'./sample_files/ddid.docx')
