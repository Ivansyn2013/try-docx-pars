import docx
import json
import time, datetime
from itertools import groupby
import re

temp_date = datetime.datetime.now()
RE_DATE = re.compile(r'\d{2}\.\d{2}\.\d{4}')


def run_config_file():
    '''создание файла для конфигурации'''
    config_dict = {'название цикла': '', 'даты проведения цикла': 'fdf', 'количество часов': ''}

    with open('config.json', 'w', encoding='utf-8') as f:
        json.dump(config_dict, f, ensure_ascii=False, indent=4)
    return None


def get_config_file():
    '''получение конфишурации из файла'''
    with open('config.json', 'r', encoding='utf-8') as f:
        config_dict = json.load(f)

    return config_dict


def time_valid(date_str):
    import re

def we_day(date):
    if date.weekday() == 1:
        day_str = 'Понедельник'
    else: day_str = '1'
    return day_str


START_DATE = datetime.datetime.strptime('06.09.2021', '%d.%m.%Y')

print(START_DATE.date())

doc = docx.Document(r'G:\Python project\try-docx-pars\sample_files\PP(572hours).docx')


raw_list = []
colum_list  = []
list1 = []
for tabel in doc.tables:
    for cell in tabel.columns[0].cells:
        raw_list.append(cell.text)


colum_list = [ el for el, _ in groupby(raw_list)]
i = 0
for tabel in doc.tables:

    for cell in tabel.columns[0].cells:
        cell.text = '!'

    for cell in tabel.columns[0].cells:
        if i >1:
            know_date = START_DATE + datetime.timedelta(days=i)
            if know_date.weekday() == 6:
                know_date = START_DATE + datetime.timedelta(days=i+1)
                cell.text = know_date.date().strftime('%d.%m.%Y') + we_day(know_date)
            else:
                cell.text = know_date.date().strftime('%d.%m.%Y') + we_day(know_date)

        i+=1



doc.save(r'G:\Python project\try-docx-pars\ready2.docx')