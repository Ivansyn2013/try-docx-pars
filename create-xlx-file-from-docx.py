import docx
import json
import time, datetime
from openpyxl import Workbook
from openpyxl.utils import *
import re

RE_LINESTART = re.compile(r'^\d{2} *')
oks_doc12 = docx.Document(r'G:\Python project\try-docx-pars\sample_files\oks_test4.docx')

# количество абзацев в документе
print(len(oks_doc12.paragraphs))

# текст первого абзаца в документе
# print(oks_doc12.paragraphs[0].text)

# текст второго абзаца в документе
# print(oks_doc12.paragraphs[6].text)
# print(oks_doc12.paragraphs[7].text)
# print(oks_doc12.paragraphs[8].text)
# print(oks_doc12.paragraphs[9].text)
# текст первого Run второго абзаца
num = 1
result = []
quest = []
ans_A = []
ans_B = []
ans_V = []
ans_G = []
ans_D = []
l = False
try_list_r = []
try_list = []
for k in oks_doc12.paragraphs:
    try_list_r.append(k.text)
try_list = list(filter(lambda x: x != '', try_list_r))
try_list = list(filter(lambda x: x != ' ', try_list))
print(try_list)
print(try_list_r)
for i in try_list:

    if RE_LINESTART.search(i) and re.search(r'^А\.', try_list[try_list.index(i) + 1]):
        quest.append(try_list[try_list.index(i)])
        ans_A.append(try_list[try_list.index(i) + 1])
        # print(i.text)
        ans_B.append(try_list[try_list.index(i) + 2])
        # print(i.text)
        ans_V.append(try_list[try_list.index(i) + 3])
        # print(i.text)
        ans_G.append(try_list[try_list.index(i) + 4])
        # print(i.text)
        ans_D.append(try_list[try_list.index(i) + 5])

result = [
    quest,
    ans_A,
    ans_B,
    ans_V,
    ans_G,
    ans_D,
]

# result.append(quest)
# result.append(ans_A)
# result.append(ans_B)
# result.append(ans_V)
# result.append(ans_G)
# result.append(ans_D)

# работа с Excel  файлом
wb = Workbook()
dest_filename = '123.xlsx'
ws1 = wb.active
ws1.title = 'questions'
for i, el in enumerate(quest):
    ws1.cell(row=i + 1, column=1).value = el
for i, el in enumerate(ans_A):
    ws1.cell(row=i + 1, column=2).value = el

for i, el in enumerate(ans_B):
    ws1.cell(row=i + 1, column=3).value = el
for i, el in enumerate(ans_V):
    ws1.cell(row=i + 1, column=4).value = el
for i, el in enumerate(ans_G):
    ws1.cell(row=i + 1, column=5).value = el
for i, el in enumerate(ans_D):
    ws1.cell(row=i + 1, column=6).value = el

wb.save(filename=dest_filename)
# if i.text != ' ':
#     print(i.text)
