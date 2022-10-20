import docx
from docx import Document
from docx.shared import RGBColor
import re
test_dict = {x:chr(x) for x in range(0,50)}
i=0
#print(test_dict)
for x in range(50,100):
    test_dict[i] = [test_dict[i],x]
    i+=1
#print(test_dict)

print(*enumerate(test_dict.values(),1))


RE_REPLASE = re.compile(r'\[([0-9,\s]+)\]')

new_dis_doc = docx.Document()
#назначени цвета тексту без docx.shared не работает
runn = new_dis_doc.add_paragraph().add_run('fdffdsfsdfsdf')
runn.font.color.rgb = RGBColor(255, 0, 0)


new_dis_doc.save(r'./sample_files/ddid.docx')

# doc_test = docx.Document()
# table = doc_test.add_table(rows = 1, cols =3)
# table.style = 'Table Grid'
#
# for num1, num2 in enumerate(test_dict.values(),1):
#
#     print(num1)
#     print(num2)
#     print(num2[0])
#     print(num2[1])
#
#
#
#     row_cells[0].text = str(num1)
#     #row_cells[1].text = str(num2[0])
#     row_cells[2].text = str(num2[1])
#
#
# doc_test.save('table.docx')
#



test_list = 'dsfdfsdfsdf [123,323,123123] fdfdffdf [123]'

m = re.finditer(RE_REPLASE,test_list)

for mat in m:
    print(mat)
    print(mat.span())

dict_rr = {1:['asda','asdasd'],2:['qwrqwe','asdfasdf']}
for i in dict_rr:
    print(i)
# print(dict_rr)
# print(len(dict_rr))
# print(dict_rr[0])