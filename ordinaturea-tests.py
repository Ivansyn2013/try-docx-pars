import docx
ord_test = docx.Document(r"C:\Users\Александр\PycharmProjects\try-docx-pars\sample_files\OKS.docx")

print(len(ord_test.paragraphs))

# текст первого абзаца в документе
print(ord_test.paragraphs[0].text)

# текст второго абзаца в документе
print(ord_test.paragraphs[1].text)

# текст первого Run второго абзаца
print(ord_test.paragraphs[1].runs[0].text)