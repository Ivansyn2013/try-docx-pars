import docx
import json
import time, datetime
from datetime import datetime as dt
import re

temp_date = datetime.datetime.now()
RE_DATE = re.compile(r'(\d{2}\.\d{2})\.\d{4}')
WEEKDAYS = ['Понедельник', 'Вторник','Среда','Четверг','Пятница', 'Суббота', 'Воскресенье']

def run_config_file():
    '''создание файла для конфигурации'''
    config_dict = {'название цикла': '', 'даты проведения цикла': 'fdf', 'количество часов': ''}

    with open('config.json', 'w', encoding='utf-8') as f:
        json.dump(config_dict, f, ensure_ascii=False, indent=4)
    return None


def get_config_file():
    '''получение конфишурации из файла'''
    with open('config.json', 'r', encoding='utf-8') as f:
        config_dict = json.load(f)

    return config_dict


def time_valid(date_str):
    import re

    RE_DATE = re.compile(r'(\d{2}\.\d{2}).\d{4}')
    if RE_DATE.search(date_str) != None:
        return True
    else:
        return False

#забор даты из конфига
START_DATE = datetime.datetime.strptime(get_config_file().get('даты проведения цикла'), '%d.%m.%Y')
# забор файла
CONFIG = get_config_file()
if CONFIG['количество часов'] == '144':
    doc = docx.Document(r'G:\Python project\try-docx-pars\input\144.docx')
    OUTFILE = r'G:\Python project\try-docx-pars\output\{}144.docx'
elif CONFIG['количество часов'] == '572':
    doc = docx.Document(r'G:\Python project\try-docx-pars\input\572.docx')
    OUTFILE = r'G:\Python project\try-docx-pars\output\{}572.docx'
else:
    raise Exception('Не указан фаил')
#забор праздничных дней из конфига
HOLYDAYS = CONFIG['праздничные дни']
HOLYDAYS = [dt.strptime(x, '%d.%m.%Y') for x in HOLYDAYS]



# выбор первого столбца таблицы

tabel = doc.tables[0]
#print(len(tabel.columns[0].cells))

k = 1
i = 1
unique, merged = set(), set()

for tabel in doc.tables:

    for cell in tabel.columns[0].cells[1:]:
        # if 'vMerge' in cell._tc.xml:
        #     flag_cellmerge = not flag_cellmerge

        delta_time = START_DATE + datetime.timedelta(days=i-1)
        tc = cell._tc # защищенный параметр с xml кодом
        cell_loc = (tc.top, tc.bottom, tc.left, tc.right) # координаты ячейки

        if delta_time.weekday() == 6 or delta_time in HOLYDAYS:   # воскресенье
            i+=1
            continue


        if cell_loc in unique:
            merged.add(cell_loc)
            continue
        else:
            unique.add(cell_loc)
            cell.text = str(str(k) + ')'+ delta_time.date().strftime('%d.%m.%Y') + '\n' + WEEKDAYS[
            delta_time.weekday()])
            k+=1
        i += 1

        print(cell.text)



# сохранение объекта файла
doc.save(OUTFILE.format(START_DATE.strftime('%d.%m.%Y')))
